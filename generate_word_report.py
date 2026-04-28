from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('Final Year Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FarmAI: An AI-Based Farm Profit Optimization System')
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph('\n' * 5)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submitted by:\nDevendra Chouhan\nPremium Farmering AI Research')
    run.font.size = Pt(14)

    doc.add_page_break()

    # --- Table of Contents Placeholder ---
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "FarmAI is a comprehensive decision-support system designed to maximize agricultural profitability "
        "through the integration of Machine Learning and real-time market data. The system provides "
        "end-to-end assistance, including crop recommendation, yield prediction, price forecasting, "
        "and profit optimization. By analyzing soil parameters (NPK) and environmental factors, "
        "FarmAI empowers farmers with data-driven insights to mitigate risks and increase net income."
    )

    doc.add_heading('2. Introduction', level=1)
    doc.add_paragraph(
        "The agricultural sector in India faces significant challenges due to unpredictable weather, "
        "fluctuating market prices, and soil degradation. Traditional farming methods often rely on "
        "intuition, which can lead to financial losses. FarmAI addresses these issues by leveraging "
        "Predictive Analytics to transform raw agricultural data into actionable intelligence."
    )

    doc.add_heading('3. System Architecture', level=1)
    doc.add_paragraph(
        "The application is built using a modular architecture that separates data processing, "
        "machine learning modeling, and user interface layers. The frontend is developed using "
        "Streamlit with a custom high-fidelity design system."
    )
    
    doc.add_heading('3.1 Module Descriptions', level=2)
    doc.add_paragraph("- Crop Recommendation Module: Random Forest Classification Model.")
    doc.add_paragraph("- Yield Prediction Module: Gradient Boosting Regression Model.")
    doc.add_paragraph("- Price Forecasting Module: ARIMA/LSTM Time-Series Analysis.")
    doc.add_paragraph("- Profit Optimization Module: Multi-variable Economic Ranking Algorithm.")

    doc.add_heading('4. Methodology', level=1)
    doc.add_paragraph(
        "The project follows a standard data science pipeline: Data Collection, Preprocessing, "
        "Feature Engineering, Model Training, and UI Integration. The core datasets were sourced "
        "from government agricultural portals, containing over 2,200 soil samples and 10 years "
        "of market price history."
    )

    doc.add_heading('5. Results and Discussion', level=1)
    doc.add_paragraph(
        "The models achieved high accuracy across all domains. The Crop Recommender achieved "
        "significant precision in matching soil types to crop requirements, while the Price Forecast "
        "module successfully identified seasonal trends across 22+ major Indian crops."
    )

    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "FarmAI successfully demonstrates the potential of AI in transforming modern agriculture. "
        "By providing a transparent and easy-to-use interface, the project brings advanced "
        "computational logic directly to the farmer's hands, ensuring a more profitable and "
        "sustainable future for the industry."
    )

    # Save the document
    file_path = "FarmAI_Project_Report.docx"
    doc.save(file_path)
    print(f"Report generated successfully: {file_path}")

if __name__ == "__main__":
    create_report()
